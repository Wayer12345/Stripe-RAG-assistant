"""Unit tests for all document parsers."""

from __future__ import annotations

import io
import json

import docx as python_docx
import pytest
from app.domain.interfaces.document_loader import RawDocument
from app.domain.models.document import DocumentProcessingStage
from app.infrastructure.parsers.csv_parser import CsvParser
from app.infrastructure.parsers.docx_parser import DocxParser
from app.infrastructure.parsers.html_parser import HtmlParser
from app.infrastructure.parsers.json_parser import JsonParser
from app.infrastructure.parsers.markdown_parser import MarkdownParser
from app.infrastructure.parsers.pdf_parser import PdfParser
from app.infrastructure.parsers.txt_parser import TxtParser
from pypdf import PdfWriter

# ---------------------------------------------------------------------------
# Byte-level fixture helpers
# ---------------------------------------------------------------------------


def _raw(
    content: bytes,
    source_type: str = "txt",
    source_path: str | None = "/data/doc.txt",
    source_name: str | None = "doc.txt",
    mime_type: str | None = None,
) -> RawDocument:
    return RawDocument(
        source_type=source_type,
        content=content,
        source_path=source_path,
        source_name=source_name,
        mime_type=mime_type,
    )


def _make_pdf_bytes(page_texts: list[str] | None = None) -> bytes:
    """Produce a multi-page PDF with blank pages (one per entry in *page_texts*)."""
    writer = PdfWriter()
    for _text in page_texts or ["Stripe payment guide"]:
        writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _make_docx_bytes(paragraphs: list[str], title: str | None = None) -> bytes:
    """Build a minimal DOCX in memory."""
    doc = python_docx.Document()
    if title:
        doc.core_properties.title = title
    for para in paragraphs:
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Common parser contract checks
# ---------------------------------------------------------------------------


def _assert_valid_document(doc, source_type: str | None = None) -> None:
    """Assert that *doc* is a valid Document with all required fields set."""
    assert doc.id.startswith("doc_")
    assert doc.text.strip()
    assert len(doc.content_hash) == 64
    assert all(c in "0123456789abcdef" for c in doc.content_hash)
    assert doc.created_at is not None
    assert doc.processing_stage == DocumentProcessingStage.PARSED
    if source_type:
        assert doc.source_type == source_type


# ---------------------------------------------------------------------------
# TxtParser
# ---------------------------------------------------------------------------


class TestTxtParser:
    def test_supports_txt(self) -> None:
        assert TxtParser().supports("txt") is True

    def test_supports_text_plain_mime(self) -> None:
        assert TxtParser().supports("other", "text/plain") is True

    def test_rejects_html(self) -> None:
        assert TxtParser().supports("html") is False

    def test_parse_valid_utf8(self) -> None:
        raw = _raw(b"Hello, Stripe!", source_type="txt")
        docs = TxtParser().parse(raw)
        assert len(docs) == 1
        assert docs[0].text == "Hello, Stripe!"

    def test_parse_invalid_utf8_fallback(self) -> None:
        raw = _raw(b"Hello \xff World", source_type="txt")
        doc = TxtParser().parse(raw)[0]
        assert "Hello" in doc.text
        assert "World" in doc.text

    def test_parse_preserves_source_metadata(self) -> None:
        raw = _raw(b"content", source_type="txt", source_path="/p/f.txt", source_name="f.txt")
        doc = TxtParser().parse(raw)[0]
        assert doc.source_path == "/p/f.txt"
        assert doc.source_name == "f.txt"

    def test_parse_deterministic_id(self) -> None:
        raw = _raw(b"same", source_type="txt")
        assert TxtParser().parse(raw)[0].id == TxtParser().parse(raw)[0].id

    def test_parse_content_hash_set(self) -> None:
        raw = _raw(b"text", source_type="txt")
        doc = TxtParser().parse(raw)[0]
        _assert_valid_document(doc, "txt")

    def test_empty_raises(self) -> None:
        raw = _raw(b"   \n", source_type="txt")
        with pytest.raises(ValueError, match="empty or whitespace"):
            TxtParser().parse(raw)

    def test_title_from_source_name(self) -> None:
        raw = _raw(b"text", source_type="txt", source_name="guide.txt")
        assert TxtParser().parse(raw)[0].title == "guide.txt"


# ---------------------------------------------------------------------------
# MarkdownParser
# ---------------------------------------------------------------------------


class TestMarkdownParser:
    def test_supports_md(self) -> None:
        assert MarkdownParser().supports("md") is True

    def test_supports_markdown(self) -> None:
        assert MarkdownParser().supports("markdown") is True

    def test_supports_mime_text_markdown(self) -> None:
        assert MarkdownParser().supports("other", "text/markdown") is True

    def test_supports_mime_x_markdown(self) -> None:
        assert MarkdownParser().supports("other", "text/x-markdown") is True

    def test_rejects_txt(self) -> None:
        assert MarkdownParser().supports("txt") is False

    def test_parse_valid_md_with_h1(self) -> None:
        raw = _raw(b"# My Title\n\nHello world.", source_type="md", source_name="doc.md")
        doc = MarkdownParser().parse(raw)[0]
        assert doc.title == "My Title"

    def test_parse_preserves_markdown_structure(self) -> None:
        content = b"# Title\n\n- item1\n- item2\n\n```python\nprint('hi')\n```"
        raw = _raw(content, source_type="md")
        doc = MarkdownParser().parse(raw)[0]
        assert "- item1" in doc.text
        assert "```python" in doc.text

    def test_title_fallback_to_source_name(self) -> None:
        raw = _raw(b"No heading here.", source_type="md", source_name="doc.md", source_path=None)
        doc = MarkdownParser().parse(raw)[0]
        assert doc.title == "doc.md"

    def test_title_fallback_to_path(self) -> None:
        raw = _raw(b"text", source_type="md", source_name=None, source_path="/path/readme.md")
        doc = MarkdownParser().parse(raw)[0]
        assert doc.title == "readme.md"

    def test_title_fallback_to_default(self) -> None:
        raw = RawDocument(source_type="md", content=b"some text")
        doc = MarkdownParser().parse(raw)[0]
        assert doc.title == "Untitled Markdown document"

    def test_empty_raises(self) -> None:
        raw = _raw(b"\n\n  ", source_type="md")
        with pytest.raises(ValueError, match="empty or whitespace"):
            MarkdownParser().parse(raw)

    def test_invalid_utf8_fallback(self) -> None:
        raw = _raw(b"# Title\n\n\xff content", source_type="md")
        doc = MarkdownParser().parse(raw)[0]
        assert "content" in doc.text

    def test_deterministic_id(self) -> None:
        raw = _raw(b"# T\ntext", source_type="md")
        assert MarkdownParser().parse(raw)[0].id == MarkdownParser().parse(raw)[0].id

    def test_preserves_source_metadata(self) -> None:
        raw = _raw(b"# T\ntext", source_type="md", source_path="/a/b.md", source_name="b.md")
        doc = MarkdownParser().parse(raw)[0]
        assert doc.source_path == "/a/b.md"
        assert doc.source_name == "b.md"

    def test_valid_document_fields(self) -> None:
        raw = _raw(b"# T\ncontent", source_type="md")
        _assert_valid_document(MarkdownParser().parse(raw)[0], "md")


# ---------------------------------------------------------------------------
# HtmlParser
# ---------------------------------------------------------------------------


class TestHtmlParser:
    def test_supports_html(self) -> None:
        assert HtmlParser().supports("html") is True

    def test_supports_htm(self) -> None:
        assert HtmlParser().supports("htm") is True

    def test_supports_mime_text_html(self) -> None:
        assert HtmlParser().supports("other", "text/html") is True

    def test_rejects_txt(self) -> None:
        assert HtmlParser().supports("txt") is False

    def test_extracts_body_text(self) -> None:
        html = b"<html><body><p>Stripe payments</p></body></html>"
        raw = _raw(html, source_type="html")
        doc = HtmlParser().parse(raw)[0]
        assert "Stripe payments" in doc.text

    def test_title_from_title_tag(self) -> None:
        html = b"<html><head><title>My Page</title></head><body><p>text</p></body></html>"
        raw = _raw(html, source_type="html")
        doc = HtmlParser().parse(raw)[0]
        assert doc.title == "My Page"

    def test_title_from_h1_when_no_title_tag(self) -> None:
        html = b"<html><body><h1>Heading</h1><p>text</p></body></html>"
        raw = _raw(html, source_type="html")
        doc = HtmlParser().parse(raw)[0]
        assert doc.title == "Heading"

    def test_title_fallback_to_source_name(self) -> None:
        html = b"<html><body><p>text</p></body></html>"
        raw = _raw(html, source_type="html", source_name="page.html")
        doc = HtmlParser().parse(raw)[0]
        assert doc.title == "page.html"

    def test_script_content_excluded(self) -> None:
        html = b"<html><body><script>alert('xss')</script><p>clean</p></body></html>"
        raw = _raw(html, source_type="html")
        doc = HtmlParser().parse(raw)[0]
        assert "alert" not in doc.text

    def test_style_content_excluded(self) -> None:
        html = b"<html><body><style>body{color:red}</style><p>visible</p></body></html>"
        raw = _raw(html, source_type="html")
        doc = HtmlParser().parse(raw)[0]
        assert "color" not in doc.text
        assert "visible" in doc.text

    def test_noscript_excluded(self) -> None:
        html = b"<html><body><noscript>enable js</noscript><p>text</p></body></html>"
        raw = _raw(html, source_type="html")
        doc = HtmlParser().parse(raw)[0]
        assert "enable js" not in doc.text

    def test_empty_body_raises(self) -> None:
        html = b"<html><body><script>x=1</script></body></html>"
        raw = _raw(html, source_type="html")
        with pytest.raises(ValueError, match="no usable text"):
            HtmlParser().parse(raw)

    def test_preserves_source_metadata(self) -> None:
        html = b"<html><body><p>content</p></body></html>"
        raw = _raw(html, source_type="html", source_path="/p/page.html", source_name="page.html")
        doc = HtmlParser().parse(raw)[0]
        assert doc.source_path == "/p/page.html"

    def test_deterministic_id(self) -> None:
        html = b"<html><body><p>x</p></body></html>"
        raw = _raw(html, source_type="html")
        assert HtmlParser().parse(raw)[0].id == HtmlParser().parse(raw)[0].id

    def test_valid_document_fields(self) -> None:
        html = b"<html><body><p>hello</p></body></html>"
        raw = _raw(html, source_type="html")
        _assert_valid_document(HtmlParser().parse(raw)[0], "html")


# ---------------------------------------------------------------------------
# PdfParser
# ---------------------------------------------------------------------------


class TestPdfParser:
    def test_supports_pdf(self) -> None:
        assert PdfParser().supports("pdf") is True

    def test_supports_mime_application_pdf(self) -> None:
        assert PdfParser().supports("other", "application/pdf") is True

    def test_rejects_txt(self) -> None:
        assert PdfParser().supports("txt") is False

    def test_parse_minimal_pdf(self) -> None:
        pdf_bytes = _make_pdf_bytes()
        raw = _raw(pdf_bytes, source_type="pdf", source_name="doc.pdf")
        docs = PdfParser().parse(raw)
        assert len(docs) == 1

    def test_page_count_in_metadata(self) -> None:
        pdf_bytes = _make_pdf_bytes(["page one", "page two"])
        raw = _raw(pdf_bytes, source_type="pdf")
        doc = PdfParser().parse(raw)[0]
        assert "page_count" in doc.metadata
        assert doc.metadata["page_count"] == 2

    def test_parser_field_in_metadata(self) -> None:
        pdf_bytes = _make_pdf_bytes()
        raw = _raw(pdf_bytes, source_type="pdf")
        doc = PdfParser().parse(raw)[0]
        assert doc.metadata.get("parser") == "pypdf"

    def test_page_markers_in_text(self) -> None:
        pdf_bytes = _make_pdf_bytes(["p1", "p2"])
        raw = _raw(pdf_bytes, source_type="pdf")
        doc = PdfParser().parse(raw)[0]
        assert "[Page 1]" in doc.text
        assert "[Page 2]" in doc.text

    def test_invalid_pdf_raises(self) -> None:
        raw = _raw(b"not a pdf at all", source_type="pdf")
        with pytest.raises(ValueError, match=r"[Cc]annot read PDF|Unexpected error"):
            PdfParser().parse(raw)

    def test_does_not_require_filesystem_path(self) -> None:
        pdf_bytes = _make_pdf_bytes()
        raw = RawDocument(source_type="pdf", content=pdf_bytes)
        docs = PdfParser().parse(raw)
        assert len(docs) == 1

    def test_title_fallback_to_source_name(self) -> None:
        pdf_bytes = _make_pdf_bytes()
        raw = _raw(pdf_bytes, source_type="pdf", source_name="report.pdf")
        doc = PdfParser().parse(raw)[0]
        assert doc.title == "report.pdf"

    def test_deterministic_id(self) -> None:
        pdf_bytes = _make_pdf_bytes()
        raw = _raw(pdf_bytes, source_type="pdf")
        assert PdfParser().parse(raw)[0].id == PdfParser().parse(raw)[0].id

    def test_valid_document_fields(self) -> None:
        pdf_bytes = _make_pdf_bytes()
        raw = _raw(pdf_bytes, source_type="pdf")
        _assert_valid_document(PdfParser().parse(raw)[0], "pdf")

    def test_preserves_source_metadata(self) -> None:
        pdf_bytes = _make_pdf_bytes()
        raw = _raw(pdf_bytes, source_type="pdf", source_path="/p/doc.pdf", source_name="doc.pdf")
        doc = PdfParser().parse(raw)[0]
        assert doc.source_path == "/p/doc.pdf"


# ---------------------------------------------------------------------------
# DocxParser
# ---------------------------------------------------------------------------


class TestDocxParser:
    def test_supports_docx(self) -> None:
        assert DocxParser().supports("docx") is True

    def test_supports_mime_docx(self) -> None:
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert DocxParser().supports("other", mime) is True

    def test_rejects_txt(self) -> None:
        assert DocxParser().supports("txt") is False

    def test_parse_minimal_docx(self) -> None:
        docx_bytes = _make_docx_bytes(["Hello DOCX"])
        raw = _raw(docx_bytes, source_type="docx", source_name="doc.docx")
        docs = DocxParser().parse(raw)
        assert len(docs) == 1

    def test_paragraph_text_extracted(self) -> None:
        docx_bytes = _make_docx_bytes(["First paragraph.", "Second paragraph."])
        raw = _raw(docx_bytes, source_type="docx")
        doc = DocxParser().parse(raw)[0]
        assert "First paragraph." in doc.text
        assert "Second paragraph." in doc.text

    def test_title_from_core_properties(self) -> None:
        docx_bytes = _make_docx_bytes(["Content"], title="Stripe Guide")
        raw = _raw(docx_bytes, source_type="docx")
        doc = DocxParser().parse(raw)[0]
        assert doc.title == "Stripe Guide"

    def test_title_fallback_to_source_name(self) -> None:
        docx_bytes = _make_docx_bytes(["content"])
        raw = _raw(docx_bytes, source_type="docx", source_name="guide.docx")
        doc = DocxParser().parse(raw)[0]
        assert doc.title == "guide.docx"

    def test_invalid_docx_raises(self) -> None:
        raw = _raw(b"not a docx", source_type="docx")
        with pytest.raises(ValueError, match=r"[Cc]annot read DOCX|Unexpected error"):
            DocxParser().parse(raw)

    def test_empty_docx_raises(self) -> None:
        docx_bytes = _make_docx_bytes([])
        raw = _raw(docx_bytes, source_type="docx")
        with pytest.raises(ValueError, match="no usable text"):
            DocxParser().parse(raw)

    def test_preserves_source_metadata(self) -> None:
        docx_bytes = _make_docx_bytes(["text"])
        raw = _raw(docx_bytes, source_type="docx", source_path="/p/f.docx", source_name="f.docx")
        doc = DocxParser().parse(raw)[0]
        assert doc.source_path == "/p/f.docx"

    def test_deterministic_id(self) -> None:
        docx_bytes = _make_docx_bytes(["same text"])
        raw = _raw(docx_bytes, source_type="docx")
        assert DocxParser().parse(raw)[0].id == DocxParser().parse(raw)[0].id

    def test_valid_document_fields(self) -> None:
        docx_bytes = _make_docx_bytes(["paragraph"])
        raw = _raw(docx_bytes, source_type="docx")
        _assert_valid_document(DocxParser().parse(raw)[0], "docx")


# ---------------------------------------------------------------------------
# JsonParser — JSON mode
# ---------------------------------------------------------------------------


class TestJsonParserJson:
    def _make_raw(self, data: object, **kwargs) -> RawDocument:
        return _raw(json.dumps(data).encode(), source_type="json", **kwargs)

    def test_supports_json(self) -> None:
        assert JsonParser().supports("json") is True

    def test_supports_mime_application_json(self) -> None:
        assert JsonParser().supports("other", "application/json") is True

    def test_rejects_txt(self) -> None:
        assert JsonParser().supports("txt", "text/plain") is False

    def test_object_with_title_and_text(self) -> None:
        raw = self._make_raw({"title": "My Doc", "text": "Hello Stripe"})
        docs = JsonParser().parse(raw)
        assert len(docs) == 1
        assert docs[0].title == "My Doc"
        assert docs[0].text == "Hello Stripe"

    def test_object_uses_content_field(self) -> None:
        raw = self._make_raw({"content": "some content"})
        doc = JsonParser().parse(raw)[0]
        assert doc.text == "some content"

    def test_object_without_text_field_pretty_prints(self) -> None:
        data = {"key": "value", "number": 42}
        raw = self._make_raw(data)
        doc = JsonParser().parse(raw)[0]
        assert '"key"' in doc.text
        assert '"value"' in doc.text

    def test_array_produces_one_doc_per_item(self) -> None:
        data = [{"title": "A", "text": "first"}, {"title": "B", "text": "second"}]
        raw = self._make_raw(data)
        docs = JsonParser().parse(raw)
        assert len(docs) == 2

    def test_array_item_index_in_metadata(self) -> None:
        data = [{"text": "a"}, {"text": "b"}]
        raw = self._make_raw(data)
        docs = JsonParser().parse(raw)
        assert docs[0].metadata["item_index"] == 0
        assert docs[1].metadata["item_index"] == 1

    def test_empty_array_raises(self) -> None:
        raw = self._make_raw([])
        with pytest.raises(ValueError, match="empty"):
            JsonParser().parse(raw)

    def test_invalid_json_raises(self) -> None:
        raw = _raw(b"{invalid json", source_type="json")
        with pytest.raises(ValueError, match="Invalid JSON"):
            JsonParser().parse(raw)

    def test_deterministic_id(self) -> None:
        raw = self._make_raw({"text": "same"})
        assert JsonParser().parse(raw)[0].id == JsonParser().parse(raw)[0].id

    def test_preserves_source_metadata_fields(self) -> None:
        raw = _raw(
            json.dumps({"text": "x"}).encode(),
            source_type="json",
            source_path="/p/f.json",
            source_name="f.json",
        )
        doc = JsonParser().parse(raw)[0]
        assert doc.source_path == "/p/f.json"
        assert doc.source_name == "f.json"

    def test_valid_document_fields(self) -> None:
        raw = self._make_raw({"title": "T", "text": "content"})
        _assert_valid_document(JsonParser().parse(raw)[0], "json")


# ---------------------------------------------------------------------------
# JsonParser — JSONL mode
# ---------------------------------------------------------------------------


class TestJsonParserJsonl:
    def _lines(self, *objects: object) -> bytes:
        return b"\n".join(json.dumps(o).encode() for o in objects)

    def test_supports_jsonl(self) -> None:
        assert JsonParser().supports("jsonl") is True

    def test_supports_mime_ndjson(self) -> None:
        assert JsonParser().supports("other", "application/x-ndjson") is True

    def test_supports_mime_jsonl(self) -> None:
        assert JsonParser().supports("other", "application/jsonl") is True

    def test_multiple_lines_produce_multiple_docs(self) -> None:
        data = self._lines({"text": "a"}, {"text": "b"}, {"text": "c"})
        raw = _raw(data, source_type="jsonl")
        docs = JsonParser().parse(raw)
        assert len(docs) == 3

    def test_empty_lines_ignored(self) -> None:
        data = b'{"text": "first"}\n\n{"text": "second"}\n'
        raw = _raw(data, source_type="jsonl")
        docs = JsonParser().parse(raw)
        assert len(docs) == 2

    def test_malformed_line_raises(self) -> None:
        data = b'{"text": "ok"}\n{bad json\n{"text": "ok2"}'
        raw = _raw(data, source_type="jsonl")
        with pytest.raises(ValueError, match="Invalid JSON on line 2"):
            JsonParser().parse(raw)

    def test_line_number_in_metadata(self) -> None:
        data = self._lines({"text": "a"}, {"text": "b"})
        raw = _raw(data, source_type="jsonl")
        docs = JsonParser().parse(raw)
        assert docs[0].metadata["line_number"] == 1
        assert docs[1].metadata["line_number"] == 2

    def test_empty_jsonl_raises(self) -> None:
        raw = _raw(b"\n\n  \n", source_type="jsonl")
        with pytest.raises(ValueError, match="no usable lines"):
            JsonParser().parse(raw)

    def test_title_extracted_from_title_field(self) -> None:
        data = b'{"title": "My Title", "text": "content"}'
        raw = _raw(data, source_type="jsonl")
        doc = JsonParser().parse(raw)[0]
        assert doc.title == "My Title"

    def test_valid_document_fields(self) -> None:
        data = b'{"title": "T", "text": "content"}'
        raw = _raw(data, source_type="jsonl")
        _assert_valid_document(JsonParser().parse(raw)[0], "jsonl")


# ---------------------------------------------------------------------------
# CsvParser
# ---------------------------------------------------------------------------


class TestCsvParser:
    def test_supports_csv(self) -> None:
        assert CsvParser().supports("csv") is True

    def test_supports_mime_text_csv(self) -> None:
        assert CsvParser().supports("other", "text/csv") is True

    def test_rejects_txt(self) -> None:
        assert CsvParser().supports("txt") is False

    def test_parse_csv_with_headers(self) -> None:
        csv_bytes = b"name,amount\nAlice,100\nBob,200\n"
        raw = _raw(csv_bytes, source_type="csv", source_name="data.csv")
        docs = CsvParser().parse(raw)
        assert len(docs) == 1

    def test_headers_in_metadata(self) -> None:
        csv_bytes = b"col_a,col_b\nv1,v2\n"
        raw = _raw(csv_bytes, source_type="csv")
        doc = CsvParser().parse(raw)[0]
        assert "headers" in doc.metadata
        assert doc.metadata["headers"] == ["col_a", "col_b"]

    def test_row_count_in_metadata(self) -> None:
        csv_bytes = b"h1,h2\nr1c1,r1c2\nr2c1,r2c2\n"
        raw = _raw(csv_bytes, source_type="csv")
        doc = CsvParser().parse(raw)[0]
        assert doc.metadata["row_count"] == 2

    def test_text_contains_header_and_values(self) -> None:
        csv_bytes = b"name,city\nAlice,London\n"
        raw = _raw(csv_bytes, source_type="csv")
        doc = CsvParser().parse(raw)[0]
        assert "name" in doc.text
        assert "Alice" in doc.text

    def test_empty_csv_raises(self) -> None:
        # RawDocument rejects empty bytes; use whitespace-only content that
        # yields no usable cells after stripping.
        raw = _raw(b"  ,  \n  ,  \n", source_type="csv")
        with pytest.raises(ValueError, match="no usable content"):
            CsvParser().parse(raw)

    def test_header_only_csv_raises(self) -> None:
        # Header row present but data row contains only empty cells.
        raw = _raw(b"col1,col2\n,\n", source_type="csv")
        with pytest.raises(ValueError, match="no usable content"):
            CsvParser().parse(raw)

    def test_preserves_source_metadata(self) -> None:
        csv_bytes = b"a,b\n1,2\n"
        raw = _raw(csv_bytes, source_type="csv", source_path="/p/f.csv", source_name="f.csv")
        doc = CsvParser().parse(raw)[0]
        assert doc.source_path == "/p/f.csv"

    def test_deterministic_id(self) -> None:
        csv_bytes = b"a,b\n1,2\n"
        raw = _raw(csv_bytes, source_type="csv")
        assert CsvParser().parse(raw)[0].id == CsvParser().parse(raw)[0].id

    def test_title_from_source_name(self) -> None:
        csv_bytes = b"x,y\n1,2\n"
        raw = _raw(csv_bytes, source_type="csv", source_name="data.csv")
        doc = CsvParser().parse(raw)[0]
        assert doc.title == "data.csv"

    def test_valid_document_fields(self) -> None:
        csv_bytes = b"col\nvalue\n"
        raw = _raw(csv_bytes, source_type="csv")
        _assert_valid_document(CsvParser().parse(raw)[0], "csv")
