"""Parser that converts raw DOCX byte payloads into normalized Document objects."""

import contextlib
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path

import docx
from docx.opc.exceptions import PackageNotFoundError

from app.domain.interfaces.document_loader import RawDocument
from app.domain.models.document import Document, DocumentProcessingStage
from app.utils.hashing import sha256_text
from app.utils.ids import make_document_id

_SUPPORTED_SOURCE_TYPES: frozenset[str] = frozenset({"docx"})
_SUPPORTED_MIME_TYPES: frozenset[str] = frozenset(
    {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
)


def _extract_table_text(table: "docx.table.Table") -> str:
    """Return a simple text rendering of a DOCX table."""
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


class DocxParser:
    """Parses raw DOCX payloads into normalized Document objects.

    Implements the :class:`~app.domain.interfaces.parser.Parser` Protocol.

    Paragraphs are extracted in document order and joined with newlines.
    Tables are linearized row-by-row with pipe-separated cells and appended
    after the paragraphs that precede them in the document body.
    """

    def supports(self, source_type: str, mime_type: str | None = None) -> bool:
        """Return ``True`` for DOCX source types and MIME types."""
        return source_type.lower() in _SUPPORTED_SOURCE_TYPES or (
            mime_type is not None and mime_type.lower() in _SUPPORTED_MIME_TYPES
        )

    def supported_source_types(self) -> set[str]:
        """Return the set of source types this parser declares support for."""
        return set(_SUPPORTED_SOURCE_TYPES)

    def parse(self, raw_document: RawDocument) -> list[Document]:
        """Parse one raw DOCX payload into a single Document.

        Args:
            raw_document: Raw payload produced by a file loader.

        Returns:
            Single-element list with the parsed Document.

        Raises:
            ValueError: If the file is not a valid DOCX or yields no text.
        """
        try:
            doc = docx.Document(BytesIO(raw_document.content))
        except PackageNotFoundError as exc:
            raise ValueError(f"Cannot read DOCX: source_path={raw_document.source_path!r}") from exc
        except Exception as exc:
            raise ValueError(
                f"Unexpected error reading DOCX: source_path={raw_document.source_path!r}"
            ) from exc

        parts: list[str] = []

        # Iterate the document body in XML order to interleave paragraphs and tables.
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if tag == "p":
                para_text = element.text_content() if hasattr(element, "text_content") else ""
                # python-docx paragraphs: match by XML element identity
                for para in doc.paragraphs:
                    if para._element is element:
                        para_text = para.text
                        break
                if para_text.strip():
                    parts.append(para_text.strip())
            elif tag == "tbl":
                for table in doc.tables:
                    if table._element is element:
                        table_text = _extract_table_text(table)
                        if table_text.strip():
                            parts.append(table_text)
                        break

        text = "\n".join(parts)
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        if not text.strip():
            raise ValueError(
                f"DOCX document yields no usable text: source_path={raw_document.source_path!r}"
            )

        content_hash = sha256_text(text)
        source_identity = raw_document.source_path or raw_document.source_name or "unknown"
        doc_id = make_document_id(source_identity, content_hash)

        # Title: core_properties.title → source_name → path filename → default
        cp_title: str | None = None
        with contextlib.suppress(Exception):
            cp_title = (doc.core_properties.title or "").strip() or None

        title = (
            cp_title
            or raw_document.source_name
            or (Path(raw_document.source_path).name if raw_document.source_path else None)
            or "Untitled DOCX document"
        )

        return [
            Document(
                id=doc_id,
                source_type=raw_document.source_type,
                source_path=raw_document.source_path,
                source_name=raw_document.source_name,
                source_mime_type=raw_document.mime_type,
                title=title,
                text=text,
                content_hash=content_hash,
                created_at=datetime.now(UTC),
                processing_stage=DocumentProcessingStage.PARSED,
                metadata=dict(raw_document.metadata),
            )
        ]
